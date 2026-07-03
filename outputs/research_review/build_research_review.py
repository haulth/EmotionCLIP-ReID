from __future__ import annotations

import html
import json
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


OUT = Path(__file__).resolve().parent
XLSX = OUT / "emotionclip_reid_32_papers.xlsx"
DOCX = OUT / "related_work_emotionclip_reid.docx"
DRAWIO = OUT / "emotionclip_reid_related_work_gap_map.drawio"
PDF_DIR = OUT / "_qa" / "pdf"
PNG_DIR = OUT / "_qa" / "pages"


PAPERS = [
    {
        "no": 1,
        "title": "Tip-Adapter: Training-free Adaption of CLIP for Few-shot Classification",
        "authors": "Renrui Zhang, Wei Zhang, Rongyao Fang, Peng Gao, Kunchang Li, Jifeng Dai, Yu Qiao, Hongsheng Li",
        "year": 2022,
        "venue_type": "Conference",
        "venue": "ECCV",
        "topic": "CLIP adaptation / adapter",
        "aim": "Thích nghi CLIP cho few-shot classification bằng cache adapter không cần huấn luyện nặng.",
        "strong": "Cho thấy CLIP có thể được adapter hóa hiệu quả; liên quan trực tiếp tới adapter/anchor trong project.",
        "weak": "Không thiết kế riêng cho FER/AU; adapter dựa trên cache feature, chưa có semantic descriptor học được.",
        "dataset": "ImageNet and 10 image classification datasets.",
        "result": "N/A - VLM adaptation; gợi ý adapter nhẹ cho FER.",
        "approach": "Cache-model adapter combines few-shot visual features with CLIP priors.",
        "link": "https://arxiv.org/abs/2207.09519",
        "group": "Foundation",
    },
    {
        "no": 2,
        "title": "Learning to Prompt for Vision-Language Models",
        "authors": "Kaiyang Zhou, Jingkang Yang, Chen Change Loy, Ziwei Liu",
        "year": 2022,
        "venue_type": "Journal",
        "venue": "IJCV",
        "topic": "Prompt learning",
        "aim": "Thay prompt thủ công bằng context token có thể học.",
        "strong": "CoOp chứng minh prompt token có thể thích nghi downstream task.",
        "weak": "Context học được có thể overfit base classes và giảm generalization.",
        "dataset": "11 image classification datasets.",
        "result": "N/A - prompt foundation; hữu ích cho emotion descriptor learning.",
        "approach": "Optimize continuous prompt context while freezing CLIP encoders.",
        "link": "https://arxiv.org/abs/2109.01134",
        "group": "Foundation",
    },
    {
        "no": 3,
        "title": "Conditional Prompt Learning for Vision-Language Models",
        "authors": "Kaiyang Zhou, Jingkang Yang, Chen Change Loy, Ziwei Liu",
        "year": 2022,
        "venue_type": "Conference",
        "venue": "CVPR",
        "topic": "Prompt learning / generalization",
        "aim": "Sinh prompt phụ thuộc từng ảnh để cải thiện generalization.",
        "strong": "CoCoOp giảm overfitting prompt cố định bằng instance-conditioned context.",
        "weak": "Tăng độ phức tạp và không tự đảm bảo diễn giải AU/cảm xúc.",
        "dataset": "Image classification base-to-new benchmarks.",
        "result": "N/A - prompt foundation; gợi ý instance-aware emotion descriptors.",
        "approach": "Meta-network generates input-conditional prompt vectors.",
        "link": "https://arxiv.org/abs/2203.05557",
        "group": "Foundation",
    },
    {
        "no": 4,
        "title": "MaPLe: Multi-Modal Prompt Learning",
        "authors": "Muhammad Uzair Khattak et al.",
        "year": 2023,
        "venue_type": "Conference",
        "venue": "CVPR",
        "topic": "Multi-modal prompt learning",
        "aim": "Học prompt ở cả nhánh vision và language để thích nghi CLIP.",
        "strong": "Liên kết prompt hai modality; phù hợp ý tưởng descriptor dẫn hướng image encoder.",
        "weak": "Thiết kế chung cho classification, chưa có ràng buộc AU/FACS.",
        "dataset": "11 classification datasets; base-to-new, cross-dataset.",
        "result": "N/A - VLM adaptation; hướng cho prompt ảnh-văn bản trong FER.",
        "approach": "Coupled visual and textual prompts in CLIP.",
        "link": "https://openaccess.thecvf.com/content/CVPR2023/html/Khattak_MaPLe_Multi-Modal_Prompt_Learning_CVPR_2023_paper.html",
        "group": "Foundation",
    },
    {
        "no": 5,
        "title": "CLIP-ReID: Exploiting Vision-Language Model for Image Re-Identification without Concrete Text Labels",
        "authors": "Siyuan Li, Li Sun, Qingli Li",
        "year": 2022,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "VLM for ReID",
        "aim": "Dùng CLIP cho ReID khi chỉ có nhãn ID, không có text thật.",
        "strong": "Khung hai giai đoạn: học text token rồi dùng text feature làm anchor cho image encoder.",
        "weak": "Token gắn với ID, ít ý nghĩa ngữ nghĩa và khó chuyển trực tiếp sang FER.",
        "dataset": "Market-1501, DukeMTMC-reID, MSMT17, VeRi, VehicleID.",
        "result": "N/A - ReID; là baseline logic cho project hiện tại.",
        "approach": "Learnable ID-specific prompt tokens + image-to-text alignment.",
        "link": "https://arxiv.org/abs/2211.13977",
        "group": "ReID",
    },
    {
        "no": 6,
        "title": "A Pedestrian is Worth One Prompt: Towards Language Guidance Person Re-Identification",
        "authors": "Shuyu Yang et al.",
        "year": 2024,
        "venue_type": "Conference",
        "venue": "CVPR",
        "topic": "Language-guided ReID",
        "aim": "Đưa mô tả thuộc tính ngôn ngữ vào ReID thay vì chỉ dựa vào ID.",
        "strong": "Cho thấy semantic attribute prompts có thể làm ReID dễ diễn giải hơn.",
        "weak": "Ngữ nghĩa pedestrian/clothing khác bản chất AU/cảm xúc tinh vi trên mặt.",
        "dataset": "Market-1501, MSMT17, CUHK03, DukeMTMC-reID.",
        "result": "N/A - ReID; hỗ trợ luận điểm dùng descriptor thay ID token.",
        "approach": "Attribute dictionary and prompt composition for person features.",
        "link": "https://openaccess.thecvf.com/content/CVPR2024/papers/Yang_A_Pedestrian_is_Worth_One_Prompt_Towards_Language_Guidance_Person_CVPR_2024_paper.pdf",
        "group": "ReID",
    },
    {
        "no": 7,
        "title": "CLIP-SCGI: Synthesized Caption-Guided Inversion for Person Re-Identification",
        "authors": "Qianru Han, Xinwei He, Zhi Liu, Sannyuya Liu, Ying Zhang, Jinhai Xiang",
        "year": 2024,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "Caption-guided ReID",
        "aim": "Tạo caption giả cho ảnh người và dùng caption dẫn hướng học ReID.",
        "strong": "Giảm phụ thuộc vào token mơ hồ bằng text mô tả giàu semantic hơn.",
        "weak": "Chất lượng caption và hallucination có thể làm lệch feature.",
        "dataset": "Four popular ReID benchmarks.",
        "result": "N/A - ReID; gợi ý dùng mô tả tự nhiên thay token ID.",
        "approach": "Pseudo captions, caption-guided inversion, cross-modal fusion.",
        "link": "https://arxiv.org/abs/2410.09382",
        "group": "ReID",
    },
    {
        "no": 8,
        "title": "CLIP-Driven Semantic Discovery Network for Visible-Infrared Person Re-Identification",
        "authors": "Xiaoyan Yu, Neng Dong, Liehuang Zhu, Hao Peng, Dapeng Tao",
        "year": 2025,
        "venue_type": "Journal",
        "venue": "IEEE Transactions on Multimedia",
        "topic": "Semantic ReID / cross-modality",
        "aim": "Dùng semantic prompt để giảm modality gap giữa visible và infrared ReID.",
        "strong": "Tách prompt theo modality và tích hợp semantic invariant vào visual features.",
        "weak": "Vẫn quanh identity matching; chưa xử lý ambiguity của nhãn cảm xúc.",
        "dataset": "SYSU-MM01, RegDB and VI-ReID benchmarks.",
        "result": "N/A - VI-ReID; hữu ích cho cross-modality semantic alignment.",
        "approach": "Modality-specific prompt learner + semantic information integration.",
        "link": "https://arxiv.org/abs/2401.05806",
        "group": "ReID",
    },
    {
        "no": 9,
        "title": "CLIPER: A Unified Vision-Language Framework for In-the-Wild Facial Expression Recognition",
        "authors": "Hanting Li, Hongjing Niu, Zhaoqing Zhu, Feng Zhao",
        "year": 2023,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "CLIP for static/dynamic FER",
        "aim": "Dùng nhiều text descriptors để biểu diễn cảm xúc tốt hơn class name.",
        "strong": "Trực tiếp ủng hộ hướng multiple expression text descriptors cho FER.",
        "weak": "Descriptor vẫn chủ yếu theo emotion label; chưa rõ AU semantic grounding.",
        "dataset": "RAF-DB, AffectNet, DFEW and FER benchmarks.",
        "result": "Reports strong FER performance across static and dynamic benchmarks.",
        "approach": "CLIP-based contrastive vision-language FER with multiple descriptors.",
        "link": "https://arxiv.org/abs/2303.00193",
        "group": "Emotion/AU",
    },
    {
        "no": 10,
        "title": "Prompting Visual-Language Models for Dynamic Facial Expression Recognition",
        "authors": "Zengqun Zhao, Ioannis Patras",
        "year": 2023,
        "venue_type": "Conference",
        "venue": "BMVC",
        "topic": "DFER-CLIP",
        "aim": "Mở rộng CLIP cho video FER bằng temporal model và text description.",
        "strong": "Text mô tả hành vi mặt tốt hơn class name; có temporal modeling.",
        "weak": "Tập trung video; phụ thuộc chất lượng mô tả LLM và fine-tuning.",
        "dataset": "DFEW, FERV39k, MAFW.",
        "result": "Reports SOTA/competitive DFER results on DFEW, FERV39k, MAFW.",
        "approach": "CLIP visual encoder + Transformer temporal encoder + learnable text token.",
        "link": "https://arxiv.org/abs/2308.13382",
        "group": "Emotion/AU",
    },
    {
        "no": 11,
        "title": "FineCLIPER: Multi-modal Fine-grained CLIP for Dynamic Facial Expression Recognition with AdaptERs",
        "authors": "Harold Chen et al.",
        "year": 2024,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "Fine-grained DFER / adapters",
        "aim": "Phân biệt biểu cảm động tinh vi bằng mô tả positive/negative và adapters.",
        "strong": "Gần project vì dùng CLIP, adapter, text supervision fine-grained.",
        "weak": "Thiết kế cho video; chưa biến AU thành semantic anchor cố định cho image encoder.",
        "dataset": "DFEW, FERV39k, MAFW.",
        "result": "Reports competitive DFER results with fine-grained CLIP adaptation.",
        "approach": "Positive/negative text descriptions, hierarchical cues, adapters.",
        "link": "https://arxiv.org/abs/2407.02157",
        "group": "Emotion/AU",
    },
    {
        "no": 12,
        "title": "Emotion-aware adaptation of CLIP model for facial expression recognition",
        "authors": "Jing Huan, Mingxing Li, Haoliang Zhou",
        "year": 2026,
        "venue_type": "Journal",
        "venue": "Artificial Intelligence Review",
        "topic": "EA-CLIP / adapter FER",
        "aim": "Tinh chỉnh CLIP cho FER bằng adapter biểu cảm và classifier giàu instance.",
        "strong": "Rất gần project: expression-aware adapter và text classifier được tăng cường.",
        "weak": "Không đặt trọng tâm vào AU descriptor có nghĩa giải phẫu.",
        "dataset": "Three in-the-wild FER benchmarks plus occlusion/pose tests.",
        "result": "Reports improved accuracy/robustness on in-the-wild FER benchmarks.",
        "approach": "Expression-aware adapter + instance-enhanced expression classifier.",
        "link": "https://link.springer.com/article/10.1007/s10462-025-11468-4",
        "group": "Emotion/AU",
    },
    {
        "no": 13,
        "title": "UA-FER: Uncertainty-aware representation learning for facial expression recognition",
        "authors": "Haoliang Zhou, Shucheng Huang, Yuqiao Xu",
        "year": 2025,
        "venue_type": "Journal",
        "venue": "Neurocomputing",
        "topic": "VLP + uncertainty FER",
        "aim": "Kết hợp CLIP/VLP và Evidential Deep Learning để giảm overconfidence trong FER.",
        "strong": "Liên kết text-visual affinity, local/global features và uncertainty calibration.",
        "weak": "Kiến trúc phức tạp; khó tái hiện đầy đủ nếu dữ liệu/protocol hạn chế.",
        "dataset": "Three in-the-wild and one in-the-lab FER benchmarks.",
        "result": "Reports SOTA improvements on multiple FER benchmarks.",
        "approach": "Multi-granularity feature decoupling + relation uncertainty calibration.",
        "link": "https://www.sciencedirect.com/science/article/pii/S0925231224020320",
        "group": "Uncertainty",
    },
    {
        "no": 14,
        "title": "Uncertain Facial Expression Recognition via Multi-task Assisted Correction",
        "authors": "Yang Liu, Xingming Zhang, Janne Kauttonen, Guoying Zhao",
        "year": 2022,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "Uncertain FER / multi-task",
        "aim": "Xử lý nhãn biểu cảm không chắc bằng AU, valence-arousal và relabeling.",
        "strong": "Dùng AU và VA làm auxiliary semantics để sửa uncertainty.",
        "weak": "Cần nhãn phụ hoặc pseudo-label đáng tin; relabeling có rủi ro lan truyền lỗi.",
        "dataset": "RAF-DB, AffectNet, Aff-Wild2.",
        "result": "Reports gains under synthetic and real uncertainty settings.",
        "approach": "Confidence weighting, AU graph branch, VA branch, feature-level relabeling.",
        "link": "https://arxiv.org/abs/2212.07144",
        "group": "Uncertainty",
    },
    {
        "no": 15,
        "title": "MAN: Mining Ambiguity and Noise for Facial Expression Recognition in the Wild",
        "authors": "Zhang et al.",
        "year": 2022,
        "venue_type": "Journal",
        "venue": "Pattern Recognition Letters",
        "topic": "Ambiguity/noise FER",
        "aim": "Tách ambiguity và noise trong FER in-the-wild để học biểu diễn bền hơn.",
        "strong": "Rất hợp phần tự phản biện: FER không chỉ sai nhãn mà còn mơ hồ tự nhiên giữa lớp cảm xúc.",
        "weak": "Không dùng VLM/AU descriptor; chủ yếu xử lý uncertainty ở không gian visual/label.",
        "dataset": "RAF-DB, AffectNet and in-the-wild FER settings.",
        "result": "Reports robustness gains by mining ambiguous/noisy samples.",
        "approach": "Ambiguity/noise mining and uncertainty-aware training.",
        "link": "https://www.sciencedirect.com/science/article/pii/S0167865522003105",
        "group": "Uncertainty",
    },
    {
        "no": 16,
        "title": "Uncertainty-Aware Label Distribution Learning for Facial Expression Recognition",
        "authors": "Nhan Le et al.",
        "year": 2023,
        "venue_type": "Conference",
        "venue": "WACV",
        "topic": "Label distribution / uncertain FER",
        "aim": "Mô hình hóa phân phối nhãn biểu cảm thay cho one-hot label cứng.",
        "strong": "Phù hợp bản chất cảm xúc mơ hồ và annotation disagreement.",
        "weak": "Không có semantic descriptor từ AU/text; cần thiết kế phân phối nhãn tốt.",
        "dataset": "RAF-DB, AffectNet and in-the-wild FER benchmarks.",
        "result": "Reports improved FER accuracy under label uncertainty.",
        "approach": "Label distribution learning with uncertainty-aware optimization.",
        "link": "https://openaccess.thecvf.com/content/WACV2023/html/Le_Uncertainty-Aware_Label_Distribution_Learning_for_Facial_Expression_Recognition_WACV_2023_paper.html",
        "group": "Uncertainty",
    },
    {
        "no": 17,
        "title": "Uncertain Label Correction via Auxiliary Action Unit Graphs for Facial Expression Recognition",
        "authors": "Yang Liu, Xingming Zhang, Guoying Zhao",
        "year": 2022,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "ULC-AG / AU graph uncertainty",
        "aim": "Sửa nhãn FER không chắc bằng auxiliary AU graph.",
        "strong": "Nối trực tiếp AU semantics với correction của nhãn emotion.",
        "weak": "Phụ thuộc chất lượng AU graph/pseudo-AU; chưa dùng VLM text descriptors.",
        "dataset": "RAF-DB, AffectNet and noisy FER settings.",
        "result": "Shows robust FER improvement under uncertain labels.",
        "approach": "Target branch with weighted regularization + AU graph branch.",
        "link": "https://arxiv.org/abs/2204.11053",
        "group": "Uncertainty",
    },
    {
        "no": 18,
        "title": "3WAUS: A Novel Three-Way Adaptive Uncertainty-Suppressing Model for Facial Expression Recognition",
        "authors": "Dong Li, Weiming Xiong, Tao Luo, Libo Zhang",
        "year": 2024,
        "venue_type": "Journal",
        "venue": "Information Sciences",
        "topic": "Uncertainty suppression / noisy labels",
        "aim": "Giảm tác động của nhãn nhiễu bằng relabeling thích nghi theo three-way decision.",
        "strong": "Bổ sung hướng xử lý uncertain/noisy labels mà không cần thêm quá nhiều nhánh phụ.",
        "weak": "Không giải quyết semantic grounding bằng text/AU; cần kết hợp với descriptor anchor nếu dùng cho project.",
        "dataset": "FER benchmarks with noisy-label/uncertainty evaluation.",
        "result": "Reports improved FER robustness under uncertainty.",
        "approach": "Three-way adaptive uncertainty suppression and dynamic relabeling.",
        "link": "https://www.sciencedirect.com/science/article/pii/S0020025524008764",
        "group": "Uncertainty",
    },
    {
        "no": 19,
        "title": "Uncertainty-Aware Multimodal Emotion Recognition through Dirichlet Parameterization",
        "authors": "Rémi Grzeczkowicz et al.",
        "year": 2026,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "Dirichlet MER / evidential fusion",
        "aim": "Hợp nhất speech, text, face bằng Dirichlet evidence và Dempster-Shafer uncertainty.",
        "strong": "Fusion không cần joint training; xử lý missing/ambiguous modalities tốt.",
        "weak": "Preprint 2026; chưa phải framework riêng cho AU semantic descriptors.",
        "dataset": "eNTERFACE05, MEAD, MELD, RAVDESS, CREMA-D.",
        "result": "Reports competitive MER accuracy with efficient uncertainty-aware fusion.",
        "approach": "Modality logits to Dirichlet evidence; uncertainty-aware fusion.",
        "link": "https://arxiv.org/abs/2602.09121",
        "group": "Multimodal",
    },
    {
        "no": 20,
        "title": "COLD Fusion: Calibrated and Ordinal Latent Distribution Fusion for Uncertainty-Aware Multimodal Emotion Recognition",
        "authors": "Mani Kumar Tellamekala et al.",
        "year": 2024,
        "venue_type": "Journal",
        "venue": "IEEE TPAMI",
        "topic": "Uncertainty-aware audiovisual MER",
        "aim": "Học uncertainty theo modality để fusion emotion audio-visual đáng tin hơn.",
        "strong": "Phân biệt độ tin cậy của từng modality; hữu ích khi face bị che hoặc audio nhiễu.",
        "weak": "Áp dụng cho audiovisual temporal context; cần giản lược nếu chỉ dùng ảnh mặt.",
        "dataset": "AVEC 2019 CES, IEMOCAP.",
        "result": "Shows benefit over model-agnostic fusion for classification/regression.",
        "approach": "Latent distributions with calibrated and ordinal variance constraints.",
        "link": "https://arxiv.org/abs/2206.05833",
        "group": "Multimodal",
    },
    {
        "no": 21,
        "title": "Hybrid Uncertainty Calibration for Multimodal Sentiment Analysis",
        "authors": "Qiuyu Pan, Zuqiang Meng",
        "year": 2024,
        "venue_type": "Journal",
        "venue": "Electronics",
        "topic": "Uncertainty calibration / multimodal fusion",
        "aim": "Hiệu chỉnh uncertainty và late fusion cho multimodal sentiment/emotion-style tasks.",
        "strong": "Gợi ý cách dùng evidential heads và uncertainty-aware fusion mà không buộc joint training phức tạp.",
        "weak": "Thuộc sentiment analysis hơn FER; chỉ nên dùng như tham khảo calibration.",
        "dataset": "Multimodal sentiment benchmarks such as MOSI/MOSEI-style settings.",
        "result": "Reports improved calibrated multimodal prediction.",
        "approach": "Evidential neural heads, hybrid uncertainty calibration, uncertain-aware late fusion.",
        "link": "https://www.mdpi.com/2079-9292/13/3/662",
        "group": "Multimodal",
    },
    {
        "no": 22,
        "title": "A Survey of Deep Learning-Based Multimodal Emotion Recognition: Speech, Text, and Face",
        "authors": "Hailun Lian, Cheng Lu, Sunan Li, Yan Zhao, Chuangao Tang, Yuan Zong",
        "year": 2023,
        "venue_type": "Journal",
        "venue": "Entropy",
        "topic": "MER survey / speech-text-face",
        "aim": "Tổng quan MER dựa trên deep learning với speech, text và face.",
        "strong": "Nguồn tốt để đặt project vào bối cảnh multimodal mà không kéo scope sang MER đầy đủ.",
        "weak": "Survey không đưa module cụ thể; cần chọn lọc ý tưởng phù hợp image-text FER.",
        "dataset": "Surveys MER datasets and fusion settings.",
        "result": "N/A - survey; dùng để phản biện scope multimodal.",
        "approach": "Taxonomy of feature extraction, fusion strategies and MER datasets.",
        "link": "https://www.mdpi.com/1099-4300/25/10/1440",
        "group": "Multimodal",
    },
    {
        "no": 23,
        "title": "Decoupled Multimodal Distilling for Emotion Recognition",
        "authors": "Yong Li, Yuanzhi Wang, Zhen Cui",
        "year": 2023,
        "venue_type": "Conference",
        "venue": "CVPR",
        "topic": "Multimodal distillation / MER",
        "aim": "Tách và chưng cất tri thức đa modality để emotion recognition linh hoạt hơn.",
        "strong": "Gợi ý học representation không phụ thuộc một modality duy nhất, phù hợp tư duy descriptor bền vững.",
        "weak": "Tập trung multimodal distillation, không phải semantic text anchor cho ảnh mặt.",
        "dataset": "IEMOCAP and multimodal emotion datasets.",
        "result": "Reports strong MER results via decoupled distillation.",
        "approach": "Flexible cross-modal knowledge distillation.",
        "link": "https://openaccess.thecvf.com/content/CVPR2023/html/Li_Decoupled_Multimodal_Distilling_for_Emotion_Recognition_CVPR_2023_paper.html",
        "group": "Multimodal",
    },
    {
        "no": 24,
        "title": "M2FNet: Multi-modal Fusion Network for Emotion Recognition in Conversation",
        "authors": "Vishal Chudasama, Purbayan Kar, Ashish Gudmalwar, Nirmesh Shah, Pankaj Wasnik, Naoyuki Onoe",
        "year": 2022,
        "venue_type": "Workshop",
        "venue": "CVPRW",
        "topic": "Multimodal fusion / ERC",
        "aim": "Kết hợp text, audio và visual cues cho emotion recognition in conversation.",
        "strong": "Có thiết kế visual branch/facial cues rõ, hữu ích khi suy nghĩ mở rộng từ image-text sang multimodal.",
        "weak": "Không phải CLIP/prompt; fusion hội thoại rộng hơn FER ảnh tĩnh.",
        "dataset": "IEMOCAP, MELD.",
        "result": "Reports SOTA weighted-F1 on MELD and IEMOCAP at publication time.",
        "approach": "Multi-head attention fusion with visual/audio/text feature extractors.",
        "link": "https://openaccess.thecvf.com/content/CVPR2022W/MULA/papers/Chudasama_M2FNet_Multi-Modal_Fusion_Network_for_Emotion_Recognition_in_Conversation_CVPRW_2022_paper.pdf",
        "group": "Multimodal",
    },
    {
        "no": 25,
        "title": "M2ER: Multimodal Emotion Recognition Based on Multi-Party Dialogue Scenarios",
        "authors": "Bo Zhang, Xiya Yang, Ge Wang, Ying Wang, Rui Sun",
        "year": 2023,
        "venue_type": "Journal",
        "venue": "Applied Sciences",
        "topic": "Multimodal dialogue emotion recognition",
        "aim": "Khai thác text, audio và video trong bối cảnh hội thoại đa người.",
        "strong": "Nhấn mạnh visual modality vẫn quan trọng trong MER, hữu ích khi phản biện mở rộng multimodal.",
        "weak": "Context hội thoại khác FER ảnh tĩnh; không cung cấp AU semantic descriptor.",
        "dataset": "MELD.",
        "result": "Reports improved multimodal emotion recognition in multi-party dialogue.",
        "approach": "Multimodal fusion scheme for text/audio/video dialogue cues.",
        "link": "https://www.mdpi.com/2076-3417/13/20/11340",
        "group": "Multimodal",
    },
    {
        "no": 26,
        "title": "RMER-DT: Robust Multimodal Emotion Recognition in Conversational Contexts Based on Diffusion and Transformers",
        "authors": "Xianxun Zhu, Yaoyang Wang, Erik Cambria, Imad Rida, José Santamaría López, Lin Cui, Rui Wang",
        "year": 2025,
        "venue_type": "Journal",
        "venue": "Information Fusion",
        "topic": "Robust MER / missing modality",
        "aim": "Tăng robustness khi modality bị thiếu/ngẫu nhiên bằng diffusion và hierarchical transformer.",
        "strong": "Rất hợp phần tự phản biện về missing/noisy modality và reliability.",
        "weak": "Quá rộng so với image-text FER; chỉ nên là hướng tham khảo tương lai.",
        "dataset": "Conversational multimodal emotion benchmarks.",
        "result": "Reports robust MER under random modality missingness.",
        "approach": "Diffusion-based reconstruction plus hierarchical transformer fusion.",
        "link": "https://www.sciencedirect.com/science/article/pii/S1566253525003410",
        "group": "Multimodal",
    },
    {
        "no": 27,
        "title": "MER-CLIP: AU-Guided Vision-Language Alignment for Micro-Expression Recognition",
        "authors": "Shifeng Liu, Xinglong Mao, Sirui Zhao, Peiming Li, Tong Xu, Enhong Chen",
        "year": 2025,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "AU-guided CLIP / MER",
        "aim": "Dùng AU/FACS description để dẫn hướng vision-language learning cho micro-expression.",
        "strong": "Rất sát thesis: AU chuyển thành text semantic cues, hỗ trợ biểu cảm tinh vi.",
        "weak": "Micro-expression/video khác static FER; cần kiểm chứng khi chuyển sang emotion classes.",
        "dataset": "Micro-expression recognition benchmarks.",
        "result": "Reports improved micro-expression recognition with AU-guided VLM.",
        "approach": "AU-guided textual descriptions aligned with visual expression features.",
        "link": "https://arxiv.org/abs/2505.05937",
        "group": "Emotion/AU",
    },
    {
        "no": 28,
        "title": "Towards End-to-End Explainable Facial Action Unit Recognition via Vision-Language Joint Learning",
        "authors": "Xuri Ge, Junchen Fu, Fuhai Chen, Shan An, Nicu Sebe, Joemon M. Jose",
        "year": 2024,
        "venue_type": "Preprint",
        "venue": "arXiv",
        "topic": "VL-FAU / explainable AU",
        "aim": "Kết hợp AU recognition với sinh mô tả ngôn ngữ local/global.",
        "strong": "Chứng minh language descriptions có thể tăng phân biệt và diễn giải AU.",
        "weak": "Tập trung AU detection; chưa ánh xạ trực tiếp AU descriptors sang emotion anchors.",
        "dataset": "DISFA, BP4D.",
        "result": "Reports superior AU recognition on most metrics.",
        "approach": "Vision-language joint learning with local AU language generation.",
        "link": "https://arxiv.org/abs/2408.00644",
        "group": "Emotion/AU",
    },
    {
        "no": 29,
        "title": "AUFormer: Vision Transformers are Parameter-Efficient Facial Action Unit Detectors",
        "authors": "Kaishen Yuan, Zitong Yu, Xin Liu, Weicheng Xie, Huanjing Yue, Jingyu Yang",
        "year": 2024,
        "venue_type": "Conference",
        "venue": "ECCV",
        "topic": "Parameter-efficient AU detection",
        "aim": "Dùng PETL và expert theo AU để phát hiện AU với ít tham số.",
        "strong": "Ủng hộ adapter/expert nhẹ cho facial muscle cues trên ViT.",
        "weak": "Không dùng text/VLM; AU detector chưa phải emotion descriptor.",
        "dataset": "BP4D, DISFA and cross-domain AU benchmarks.",
        "result": "Reports SOTA robust AU detection without extra relevant data.",
        "approach": "MoKE experts injected into frozen ViT + difficulty-aware asymmetric loss.",
        "link": "https://arxiv.org/abs/2403.04697",
        "group": "Emotion/AU",
    },
    {
        "no": 30,
        "title": "Norface: Improving Facial Expression Analysis by Identity Normalization",
        "authors": "Hanwei Liu, Rudong An, Zhimeng Zhang, Bowen Ma, Wei Zhang, Yan Song, Yujing Hu, Wei Chen, Yu Ding",
        "year": 2024,
        "venue_type": "Conference",
        "venue": "ECCV",
        "topic": "Identity normalization / FEA",
        "aim": "Loại nhiễu identity, pose, background để giữ expression-consistent signal.",
        "strong": "Rất quan trọng cho robust FER/AU vì tách expression khỏi identity noise.",
        "weak": "Cần normalization network và data pipeline riêng; không tận dụng text semantics.",
        "dataset": "AU detection, AU intensity, FER and cross-dataset benchmarks.",
        "result": "Reports SOTA in AU detection/intensity and FER settings.",
        "approach": "Normalize faces to common identity + Mixture of Experts classifier.",
        "link": "https://arxiv.org/abs/2407.15617",
        "group": "Emotion/AU",
    },
    {
        "no": 31,
        "title": "Learning Contrastive Feature Representations for Facial Action Unit Detection",
        "authors": "Ziqiao Shang et al.",
        "year": 2025,
        "venue_type": "Journal",
        "venue": "Pattern Recognition",
        "topic": "AUNCE / contrastive AU",
        "aim": "Học feature AU phân biệt bằng contrastive learning chống imbalance và noisy labels.",
        "strong": "Phù hợp project vì AU rất dễ mất cân bằng và nhãn nhiễu.",
        "weak": "Không có language descriptor; cần ghép thêm text/AU semantics nếu dùng với CLIP.",
        "dataset": "BP4D, DISFA, BP4D+, GFT, Aff-Wild2.",
        "result": "Reports superior AU detection across five datasets.",
        "approach": "AUNCE contrastive loss with negative reweighting and positive sampling.",
        "link": "https://www.sciencedirect.com/science/article/abs/pii/S0031320325014098",
        "group": "Emotion/AU",
    },
    {
        "no": 32,
        "title": "A Comprehensive Review of Multimodal Emotion Recognition: Techniques, Challenges, and Future Directions",
        "authors": "You Wu, Qingwei Mi, Tianhan Gao",
        "year": 2025,
        "venue_type": "Journal",
        "venue": "Biomimetics",
        "topic": "MER survey",
        "aim": "Tổng quan tiến bộ và thách thức trong multimodal emotion recognition.",
        "strong": "Hữu ích để đặt project vào bối cảnh multimodal, fusion, uncertainty và robustness.",
        "weak": "Survey không đưa kiến trúc cụ thể; cần chọn lọc bài liên quan trực tiếp.",
        "dataset": "Covers MER datasets such as MELD, IEMOCAP, CMU-MOSEI and others.",
        "result": "N/A - survey; dùng làm bản đồ nghiên cứu.",
        "approach": "Taxonomy of multimodal emotion features, fusion, datasets and challenges.",
        "link": "https://www.mdpi.com/2313-7673/10/7/418",
        "group": "Multimodal",
    },
]


SYNTHESIS = [
    ("CLIP/VLM adaptation", "Tip-Adapter, CoOp, CoCoOp, MaPLe", "VLM có semantic space mạnh và có thể thích nghi bằng prompt/adapter.", "Prompt chung tạo mô tả cấp lớp, chưa chạm tới tín hiệu vi mô của cơ mặt.", "Kế thừa prompt/adapter learning nhưng ràng buộc vào emotion/AU semantic descriptors."),
    ("ReID with language guidance", "CLIP-ReID, Pedestrian Prompt, CLIP-SCGI, CSDN", "Text/prompt có thể anchor visual encoder trong ReID.", "CLIP-ReID học ID-specific tokens: phân biệt tốt nhưng thiếu nghĩa và không giải thích được biểu cảm.", "Giữ khung hai giai đoạn, thay ID token bằng descriptor cảm xúc/AU có thể diễn giải."),
    ("CLIP-based FER/DFER", "CLIPER, DFER-CLIP, FineCLIPER, EA-CLIP", "Text descriptors và adapters cải thiện FER/DFER.", "Nhiều mô tả vẫn quanh class name hoặc video context, chưa biến AU thành anchor ổn định.", "Học descriptor theo emotion + AU, dùng global/local image-text alignment cho ảnh mặt."),
    ("AU/FACS semantics", "MER-CLIP, VL-FAU, AUFormer, Norface, AUNCE", "AU là cầu nối giải phẫu giữa ảnh mặt và ngôn ngữ.", "AU label thiếu, noisy hoặc domain-specific; dùng AU cứng có thể làm mô hình nhiễu.", "Dùng AU như prior có điều kiện/pseudo-AU có kiểm soát; bắt buộc có ablation."),
    ("Uncertainty-aware FER", "MTAC, MAN, LDL, ULC-AG, UA-FER, 3WAUS", "FER có label ambiguity và overconfidence.", "Uncertainty có thể làm mục tiêu nghiên cứu bị loãng nếu đưa vào trước khi descriptor baseline vững.", "Đặt uncertainty là nhánh hiệu chỉnh/calibration sau semantic alignment."),
    ("Multimodal MER", "Dirichlet MER, COLD Fusion, HUC, DMD, M2FNet, M2ER, RMER-DT", "Fusion text/audio/vision tăng robustness và xử lý missing/noisy modality.", "Thêm audio/video ngay sẽ đổi bài toán và tăng chi phí dữ liệu.", "Chỉ mượn nguyên lý reliability/fusion; project v1 tập trung image-text descriptors."),
]


THESIS_ROWS = [
    ("Vấn đề lõi", "CLIP-ReID học token theo từng ID để phân biệt identity. Token này không phải mô tả tự nhiên, nên không phù hợp nếu mục tiêu là nhận diện biểu cảm có diễn giải."),
    ("Tiêu chí chọn paper", "Danh sách chính chỉ giữ các công trình từ năm 2022 trở lại đây. CLIP 2021 được xem là tiền đề kỹ thuật, không tính vào bảng 32 paper."),
    ("Khoảng trống nghiên cứu", "Các hướng CLIP-FER đã dùng text prompt/adapters, còn AU/uncertainty/multimodal xử lý từng mảnh vấn đề. Chưa có khung gọn thay ID token bằng emotion/AU descriptors làm semantic anchors cho image encoder."),
    ("Giả thuyết nghiên cứu", "Nếu descriptor được học theo emotion/AU semantics và cố định như text anchors, visual encoder sẽ học feature bám vào tín hiệu biểu cảm thay vì chỉ tối ưu classifier một nhãn."),
    ("Đóng góp chính nên nhấn mạnh", "Chuyển cơ chế Stage 1/Stage 2 của CLIP-ReID từ ID-specific prompt learning sang emotion/AU semantic descriptor learning cho FER."),
    ("Đóng góp phụ", "Expression-aware adapters, local patch-to-text alignment và uncertainty calibration là thành phần hỗ trợ để tăng robustness, không phải luận điểm trung tâm."),
    ("Phạm vi nên tránh tuyên bố quá mức", "Không claim full multimodal MER hoặc full AU detection SOTA nếu chưa có audio/video/AU protocol. Multimodal và uncertainty chỉ là nguồn cảm hứng và hướng mở rộng."),
]


SELF_CRITIQUE = [
    ("Related work quá rộng", "Thêm multimodal/uncertainty làm tài liệu đầy đủ hơn nhưng dễ làm mờ thesis.", "Trong Word và Excel, gom chúng thành nhánh hỗ trợ robustness; luôn quay về descriptor anchor."),
    ("Emotion label có sẵn, sao cần prompt?", "Happy/sad/neutral là nhãn ngữ nghĩa nhưng quá thô; class-name prompt không mô tả cơ chế mặt.", "Nhấn mạnh AU/FACS descriptor cung cấp tín hiệu local muscle cues và khả năng diễn giải."),
    ("AU label không phải dataset nào cũng có", "Nếu bắt buộc AU thật, phương pháp khó áp dụng rộng.", "Đặt AU là prior có điều kiện: v1 dùng emotion prompt, sau đó ablation AU/pseudo-AU."),
    ("Uncertainty có thể phức tạp hóa mô hình", "EDL/Dirichlet dễ làm training khó và che mất đóng góp chính.", "Chỉ dùng sau khi baseline semantic alignment ổn; báo cáo như calibration module."),
    ("So với EA-CLIP/FineCLIPER khác gì?", "Các paper này đã có adapter/descriptor cho FER/DFER.", "Điểm khác cần nêu rõ: kế thừa CLIP-ReID hai giai đoạn và thay ID token bằng emotion/AU anchors cố định."),
    ("Đánh giá phải chứng minh descriptor thật sự hữu ích", "Nếu chỉ tăng accuracy nhỏ, đóng góp có thể bị xem là prompt tuning bình thường.", "Bổ sung ablation: class prompt vs learned descriptor vs AU descriptor; kiểm tra confusion, occlusion/pose subset, feature visualization."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "20 Papers" if False else "32 Papers"

    headers = [
        "No",
        "Article name/Tên bài báo",
        "Author(s)/DS tác giả",
        "Year Pub/năm công bố",
        "Type/Paper conf or Journal",
        "Journal name/Tên tạp chí",
        "Type/loại/chủ đề NC",
        "Aim of the paper/Mục tiêu của bài báo",
        "Strong point/Điểm mạnh",
        "Weakness/Hạn chế",
        "Dataset/Bộ dữ liệu",
        "Result in Emotion",
        "Approach/Tiếp cận để giải quyết",
        "Link",
    ]
    ws.append(headers)
    for p in PAPERS:
        row = [
            p["no"],
            p["title"],
            p["authors"],
            p["year"],
            p["venue_type"],
            p["venue"],
            p["topic"],
            p["aim"],
            p["strong"],
            p["weak"],
            p["dataset"],
            p["result"],
            p["approach"],
            p["link"],
        ]
        ws.append(row)
        ws.cell(row=ws.max_row, column=14).hyperlink = p["link"]
        ws.cell(row=ws.max_row, column=14).style = "Hyperlink"

    header_fill = PatternFill("solid", fgColor="1F4E78")
    group_fills = {
        "Foundation": "EAF2F8",
        "ReID": "E2F0D9",
        "Emotion/AU": "FFF2CC",
        "Uncertainty": "FCE4D6",
        "Multimodal": "EDE7F6",
    }
    thin = Side(style="thin", color="D9E2F3")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
    for r in range(2, ws.max_row + 1):
        group = PAPERS[r - 2]["group"]
        fill = PatternFill("solid", fgColor=group_fills[group])
        for c in range(1, len(headers) + 1):
            cell = ws.cell(r, c)
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            cell.font = Font(size=9)
        ws.cell(r, 1).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(r, 4).alignment = Alignment(horizontal="center", vertical="center")

    widths = [6, 42, 32, 10, 18, 24, 24, 38, 38, 38, 32, 38, 38, 44]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.row_dimensions[1].height = 42
    for r in range(2, ws.max_row + 1):
        ws.row_dimensions[r].height = 84
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    table = Table(displayName="PaperReviewTable", ref=ws.dimensions)
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=False, showColumnStripes=False)
    ws.add_table(table)

    sm = wb.create_sheet("Synthesis Matrix")
    sm.append(["Nhóm nghiên cứu", "Paper tiêu biểu", "Điểm mạnh kế thừa", "Research gap", "Ý nghĩa cho EmotionCLIP-ReID"])
    for row in SYNTHESIS:
        sm.append(row)
    for cell in sm[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sm.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            cell.font = Font(size=10)
    for idx, width in enumerate([28, 40, 48, 48, 50], start=1):
        sm.column_dimensions[get_column_letter(idx)].width = width
    for r in range(2, sm.max_row + 1):
        sm.row_dimensions[r].height = 72
    sm.freeze_panes = "A2"

    thesis = wb.create_sheet("Thesis Focus")
    thesis.append(["Mục", "Nội dung nhấn mạnh"])
    for row in THESIS_ROWS:
        thesis.append(row)
    for cell in thesis[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")
    for row in thesis.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            cell.font = Font(size=11)
    thesis.column_dimensions["A"].width = 30
    thesis.column_dimensions["B"].width = 115
    for r in range(2, thesis.max_row + 1):
        thesis.row_dimensions[r].height = 58

    critique = wb.create_sheet("Self-Critique")
    critique.append(["Điểm cần phản biện", "Rủi ro học thuật", "Cách chỉnh luận điểm/phương pháp"])
    for row in SELF_CRITIQUE:
        critique.append(row)
    for cell in critique[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in critique.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            cell.font = Font(size=10)
    for idx, width in enumerate([32, 52, 62], start=1):
        critique.column_dimensions[get_column_letter(idx)].width = width
    for r in range(2, critique.max_row + 1):
        critique.row_dimensions[r].height = 68
    critique.freeze_panes = "A2"

    gap = wb.create_sheet("Research Gap")
    gap_rows = [
        ("Problem", "CLIP-ReID giải quyết ReID bằng token theo ID. Khi chuyển sang FER, token theo ID không trả lời câu hỏi 'mặt đang biểu hiện cảm xúc gì và dựa trên dấu hiệu cơ mặt nào'."),
        ("Gap 1", "FER cần descriptor có nghĩa về emotion/AU, không chỉ class name như happy/sad và cũng không phải ID token không diễn giải được."),
        ("Gap 2", "Các phương pháp CLIP-FER/adapters cải thiện biểu diễn nhưng thường chưa đặt descriptor AU/FACS làm anchor cố định dẫn hướng image encoder."),
        ("Gap 3", "AU/FACS rất có giá trị nhưng nhãn AU không luôn sẵn có; vì vậy cần thiết kế AU như prior có điều kiện và đánh giá bằng ablation."),
        ("Gap 4", "Uncertainty/multimodal giúp robustness nhưng không nên làm lệch trọng tâm. Chúng hỗ trợ calibration, còn đóng góp chính là semantic descriptor learning."),
        ("Proposed Direction", "Stage 1 học emotion/AU semantic descriptors bằng CLIP text encoder; Stage 2 cố định descriptors làm anchors để fine-tune image encoder/adapters với global-local alignment và optional uncertainty calibration."),
    ]
    gap.append(["Mục", "Nội dung"])
    for row in gap_rows:
        gap.append(row)
    for cell in gap[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")
    for row in gap.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            cell.font = Font(size=11)
    gap.column_dimensions["A"].width = 24
    gap.column_dimensions["B"].width = 110
    for r in range(2, gap.max_row + 1):
        gap.row_dimensions[r].height = 52

    wb.save(XLSX)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return p


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Công Trình Liên Quan Cho EmotionCLIP-ReID")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Tự phản biện related work và làm rõ hướng nghiên cứu: emotion/AU descriptors as semantic anchors")
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Arial"

    add_heading(doc, "1. Luận điểm tổng quan", 1)
    add_body(
        doc,
        "Các công trình gần đây cho thấy CLIP và prompt learning có thể đưa tri thức ngôn ngữ vào mô hình thị giác. "
        "Tuy nhiên, CLIP-ReID gốc học token theo từng ID, nên text side chủ yếu là embedding phân biệt định danh chứ chưa phải mô tả có nghĩa. "
        "Khi chuyển sang biểu cảm mặt, vấn đề không còn là phân biệt 'người/xe nào', mà là học feature bám vào dấu hiệu biểu cảm có nghĩa: emotion category, AU/FACS và vùng cơ mặt liên quan. "
        "Vì vậy hướng nghiên cứu hợp lý hơn là học emotion/AU semantic descriptors để đóng vai trò semantic anchor cho image encoder.",
    )
    add_body(
        doc,
        "Bộ related work này được mở rộng theo sáu nhánh: VLM/prompt learning, ReID có hướng dẫn ngôn ngữ, CLIP-based FER/DFER, AU/FACS semantic guidance, uncertainty-aware FER, và multimodal emotion recognition. "
        "Tự phản biện lại, sáu nhánh này không nên được trình bày như sáu đóng góp ngang nhau. Luận điểm chính chỉ có một: thay ID-specific prompt learning của CLIP-ReID bằng emotion/AU descriptor learning. "
        "Uncertainty, multimodal fusion và adapter là các nhánh hỗ trợ để tăng độ bền và kiểm chứng, không phải phần thay thế luận điểm trung tâm.",
    )
    add_body(
        doc,
        "Theo yêu cầu cập nhật, danh sách paper chính chỉ giữ các công trình từ năm 2022 trở lại đây. CLIP 2021 vẫn là tiền đề kỹ thuật của toàn bộ hướng vision-language, nhưng không được tính như một dòng paper trong bảng để tránh lệch tiêu chí thời gian.",
    )
    add_body(
        doc,
        "Câu hỏi nghiên cứu nên viết rõ: liệu descriptor có nghĩa theo emotion/AU, khi được học và cố định như text anchors, có dẫn hướng visual encoder học biểu diễn biểu cảm tốt hơn class-name prompt hoặc ID-like tokens hay không?",
    )

    sections = [
        (
            "2. Vision-language foundation and prompt learning",
            "Tip-Adapter, CoOp, CoCoOp và MaPLe cho thấy CLIP/VLM có thể được thích nghi bằng adapter hoặc prompt học được thay vì chỉ dùng prompt viết tay. "
            "Điểm mạnh của nhóm này là khả năng chuyển tri thức ngôn ngữ sang downstream tasks. Tự phản biện: nếu chỉ dùng prompt dạng 'a happy face', project sẽ chỉ là prompt tuning thông thường. "
            "Vì vậy cần nhấn mạnh descriptor có cấu trúc theo emotion/AU, không dừng ở class-name prompt.",
        ),
        (
            "3. From CLIP-ReID ID tokens to semantic descriptors",
            "CLIP-ReID chứng minh text feature có thể làm anchor cho image encoder trong ReID, nhưng token được học theo ID nên khó diễn giải. "
            "Các hướng như Pedestrian Prompt, CLIP-SCGI và CSDN cố gắng đưa semantic caption hoặc attribute prompt vào ReID. Điểm cần nói sắc hơn: project không phủ nhận CLIP-ReID, mà kế thừa đúng phần mạnh nhất của nó là cơ chế Stage 1/Stage 2. "
            "Phần thay đổi cốt lõi là chuyển đơn vị học từ ID token sang emotion/AU semantic descriptor.",
        ),
        (
            "4. Facial expression recognition with CLIP and dynamic emotion prompts",
            "CLIPER, DFER-CLIP, FineCLIPER và EA-CLIP cho thấy mô tả văn bản, temporal modeling và adapter giúp FER/DFER hưởng lợi từ VLM. "
            "Tuy vậy nhiều phương pháp vẫn dùng mô tả cảm xúc tổng quát hoặc phụ thuộc video. Tự phản biện với EA-CLIP/FineCLIPER: nếu project chỉ thêm adapter thì đóng góp yếu. "
            "Do đó cần đặt trọng tâm vào descriptor pipeline: học descriptor, cố định descriptor, rồi dùng descriptor kéo visual feature bằng global/local image-text alignment.",
        ),
        (
            "5. AU/FACS as interpretable semantic descriptors",
            "MER-CLIP, VL-FAU, AUFormer, Norface và AUNCE nhấn mạnh AU/FACS như lớp trung gian giữa chuyển động cơ mặt và cảm xúc. "
            "AU giúp descriptor không chỉ là tên lớp như 'happy' hay 'sad', mà có thể mô tả cheek raiser, lip corner puller, brow lowerer. Điểm cần thận trọng: không phải dataset nào cũng có AU. "
            "Vì vậy AU không nên được viết như giả định bắt buộc; nó là prior có điều kiện hoặc pseudo-AU cần kiểm tra bằng ablation.",
        ),
        (
            "6. Uncertainty-aware FER and label correction",
            "MTAC, MAN, Uncertainty-aware Label Distribution Learning, ULC-AG, UA-FER và 3WAUS đều xem ambiguity là bản chất của FER in-the-wild. "
            "Các paper này chỉ ra softmax dễ overconfident khi ảnh bị che, pose lệch hoặc label mơ hồ. Tuy nhiên uncertainty không phải đóng góp chính của project v1. "
            "Nên trình bày nó như cơ chế calibration sau khi semantic descriptor alignment đã ổn, tránh làm proposal trở thành một mô hình quá nhiều module.",
        ),
        (
            "7. Multimodal emotion recognition and lessons for robust descriptor learning",
            "DMD, M2FNet, M2ER, RMER-DT, COLD Fusion, Hybrid Uncertainty Calibration và Dirichlet MER cho thấy emotion recognition mạnh hơn khi biết cân nhắc modality đáng tin. "
            "Tự phản biện: đưa multimodal vào quá nhiều sẽ làm lệch bài toán từ image-text FER sang audio-video-text MER. "
            "Vì vậy nhóm paper này chỉ nên dùng để tham khảo nguyên lý reliability, missing-modality robustness và uncertainty fusion; v1 của project vẫn nên là image-text descriptor learning.",
        ),
    ]
    for heading, text in sections:
        add_heading(doc, heading, 1)
        add_body(doc, text)

    add_heading(doc, "8. Dẫn dắt về phương pháp hiện tại", 1)
    add_body(
        doc,
        "Từ các nhóm công trình trên, hướng nghiên cứu của project có thể được phát biểu rõ: thay vì học token mô tả theo ID như CLIP-ReID, mô hình học emotion/AU semantic descriptors bằng CLIP text encoder. "
        "Các descriptor này được cố định sau Stage 1 và dùng làm semantic anchors trong Stage 2 để kéo visual embedding về vùng cảm xúc đúng. Đây là câu chuyện chính cần giữ xuyên suốt khi viết công trình liên quan.",
    )
    add_body(
        doc,
        "Phiên bản nên triển khai có kiểm soát gồm: EmotionPromptLearner, CLIP text encoder, descriptor cố định, image encoder với expression-aware adapters, global-local image-text alignment, classifier head và uncertainty/evidence output. "
        "AU/FACS và uncertainty không nên được tuyên bố là điều kiện bắt buộc ngay từ đầu; chúng nên là các ablation để chứng minh đóng góp. Cách viết này giúp proposal không bị phản biện là ghép nhiều paper, mà là một chuyển đổi có chủ đích từ CLIP-ReID sang FER.",
    )

    doc.add_page_break()
    add_heading(doc, "9. Tự phản biện và cách chỉnh luận điểm", 1)
    add_body(
        doc,
        "Bảng dưới đây dùng để tự kiểm tra khi viết luận văn/báo cáo: mỗi hướng liên quan phải phục vụ luận điểm descriptor anchor. Nếu một paper không giúp trả lời câu hỏi này, chỉ nên để ở phần tham khảo phụ.",
    )
    critique_table = doc.add_table(rows=1, cols=3)
    critique_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    critique_table.style = "Table Grid"
    crit_hdr = critique_table.rows[0].cells
    for i, h in enumerate(["Điểm cần phản biện", "Rủi ro học thuật", "Cách chỉnh"]):
        set_cell_text(crit_hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(crit_hdr[i], "1F4E78")
    for row in SELF_CRITIQUE:
        cells = critique_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if i == 0:
                set_cell_shading(cells[i], "EAF2F8")

    add_heading(doc, "10. Bảng tổng hợp nhóm paper", 1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Nhóm", "Paper tiêu biểu", "Điểm mạnh", "Gap", "Ý nghĩa cho project"]):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E78")
    for row in SYNTHESIS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if i == 0:
                set_cell_shading(cells[i], "EAF2F8")

    add_heading(doc, "11. Paper nên ưu tiên đọc sâu", 1)
    priority = [
        "CLIP-ReID: hiểu khung Stage 1/Stage 2 và vai trò text anchor.",
        "EA-CLIP và FineCLIPER: học adapter và descriptor fine-grained cho FER/DFER.",
        "MER-CLIP và VL-FAU: chuyển AU/FACS thành mô tả ngôn ngữ có diễn giải.",
        "MTAC, UA-FER và Dirichlet MER: thiết kế uncertainty/correction sao cho không làm mô hình quá phức tạp.",
        "DMD, M2FNet, COLD Fusion: tham khảo fusion/uncertainty nếu mở rộng sang multimodal.",
    ]
    for item in priority:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.add_run(item).font.size = Pt(10.5)

    add_heading(doc, "12. Nguồn tham khảo chính", 1)
    for p in PAPERS:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(f"[{p['no']}] {p['authors']} ({p['year']}). {p['title']}. {p['venue']}. ")
        run.font.size = Pt(8.5)
        run.font.name = "Arial"
        add_hyperlink(para, "Link", p["link"])

    doc.save(DOCX)


def mx_cell(id_, value, x, y, w, h, fill, stroke="#64748B", font=14):
    value = html.escape(value).replace("\n", "&lt;br&gt;")
    return f'''<mxCell id="{id_}" value="{value}" style="rounded=1;whiteSpace=wrap;html=1;fillColor={fill};strokeColor={stroke};strokeWidth=2;arcSize=8;fontSize={font};fontFamily=Arial;fontColor=#111827;" vertex="1" parent="1"><mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'''


def mx_edge(id_, source, target, label=""):
    label = html.escape(label)
    return f'''<mxCell id="{id_}" value="{label}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;strokeColor=#475569;strokeWidth=2;endArrow=block;endFill=1;fontSize=12;fontFamily=Arial;" edge="1" parent="1" source="{source}" target="{target}"><mxGeometry relative="1" as="geometry"/></mxCell>'''


def build_drawio():
    cells = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>']
    cells += [
        mx_cell("title", "Related Work Reasoning Map\nCore thesis: replace ID-specific tokens with emotion/AU semantic anchors", 285, 20, 700, 75, "#EAF2F8", "#1F4E78", 19),
        mx_cell("foundation", "VLM Adaptation (2022+)\nTip-Adapter, CoOp, CoCoOp, MaPLe\nStrong: prompt/adapter learning\nRisk: class-name prompt is too coarse", 40, 140, 260, 145, "#EAF2F8", "#1F4E78"),
        mx_cell("reid", "Language-Guided ReID\nCLIP-ReID, Pedestrian Prompt, CLIP-SCGI, CSDN\nKeep: 2-stage text-anchor training\nChange: ID token -> emotion/AU descriptor", 350, 140, 280, 145, "#E2F0D9", "#548235"),
        mx_cell("emotion", "CLIP-based FER/DFER\nCLIPER, DFER-CLIP, FineCLIPER, EA-CLIP\nRisk: adapter-only contribution is weak\nNeed: descriptor pipeline", 700, 140, 300, 145, "#FFF2CC", "#C65911"),
        mx_cell("au", "AU/FACS Semantics\nMER-CLIP, VL-FAU, AUFormer, Norface, AUNCE\nStrong: interpretable muscle cues\nGap: AU labels are missing/noisy", 1060, 140, 300, 145, "#FCE4D6", "#C00000"),
        mx_cell("unc", "Uncertainty FER (2022+)\nMTAC, MAN, LDL, ULC-AG, UA-FER, 3WAUS\nUse as calibration after baseline\nDo not let it become the main thesis", 180, 380, 310, 145, "#F4CCCC", "#A61C00"),
        mx_cell("multi", "Multimodal MER (2022+)\nDirichlet MER, COLD Fusion, HUC, DMD, M2FNet\nBorrow reliability/fusion ideas\nV1 remains image-text", 560, 380, 340, 145, "#EDE7F6", "#674EA7"),
        mx_cell("gap", "Research Gap\nNo compact framework yet converts CLIP-ReID's ID-token learning into emotion/AU descriptor anchors for FER.\nNeed interpretability + robustness without over-scoping.", 960, 380, 340, 145, "#F8FAFC", "#334155", 15),
        mx_cell("project", "Current Project: EmotionCLIP-ReID\nStage 1: learn emotion/AU descriptors\nStage 2: freeze anchors and guide image encoder/adapters\nAuxiliary: local alignment + uncertainty ablations", 345, 640, 690, 145, "#D9EAD3", "#38761D", 17),
    ]
    cells += [
        mx_edge("e1", "foundation", "reid", "prompt learning"),
        mx_edge("e2", "reid", "emotion", "replace ID tokens"),
        mx_edge("e3", "emotion", "au", "fine-grained semantics"),
        mx_edge("e4", "au", "gap", "interpretability gap"),
        mx_edge("e5", "unc", "gap", "ambiguity"),
        mx_edge("e6", "multi", "gap", "robustness"),
        mx_edge("e7", "gap", "project", "proposed answer"),
        mx_edge("e8", "emotion", "project", "adapters + descriptors"),
        mx_edge("e9", "unc", "project", "calibration"),
        mx_edge("e10", "multi", "project", "future extension"),
    ]
    xml = f'''<mxfile host="app.diagrams.net" modified="2026-05-12T00:00:00.000Z" agent="Codex" version="24.7.17">
  <diagram id="emotionclip-related-work-gap-map" name="Related Work Gap Map">
    <mxGraphModel dx="1480" dy="900" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1400" pageHeight="860" math="0" shadow="0">
      <root>
        {''.join(cells)}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>'''
    DRAWIO.write_text(xml, encoding="utf-8")


def verify_docx_render():
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if not soffice.exists():
        return False, "LibreOffice not found"
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(DOCX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf = PDF_DIR / f"{DOCX.stem}.pdf"
    if not pdf.exists():
        return False, "PDF not produced"
    subprocess.run(
        ["pdftoppm", "-png", "-r", "120", str(pdf), str(PNG_DIR / "page")],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return True, "Rendered DOCX to PDF/PNG"


def main():
    build_excel()
    build_docx()
    build_drawio()
    ok, msg = verify_docx_render()
    summary = {
        "xlsx": str(XLSX),
        "docx": str(DOCX),
        "drawio": str(DRAWIO),
        "docx_render_ok": ok,
        "docx_render_message": msg,
        "paper_count": len(PAPERS),
    }
    (OUT / "build_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
